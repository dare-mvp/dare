from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/legal/DARE_Nigeria_Legal_and_Licensing_Brief.docx")

BLUE = RGBColor(31, 77, 120)
MID_BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(20, 20, 24)
MUTED = RGBColor(89, 89, 89)
ORANGE = RGBColor(255, 85, 0)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F8F9FB"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_dxa):
                set_fixed_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_run(paragraph, text: str, bold=False, italic=False, color=None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="DADCE0", size="6")
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.style = "Normal"
    add_run(p, title + " ", bold=True, color=BLUE)
    add_run(p, body)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table, color="DADCE0", size="6")
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_fill(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = DARK
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.style = "Normal"
            p.add_run(value)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 16, 8),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("DARE legal and licensing brief - prepared for counsel")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build_document() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("DARE")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = ORANGE

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    add_run(
        subtitle,
        "Explanatory Brief for Nigerian Business Formation, Licensing, and Regulatory Review",
        bold=True,
        color=BLUE,
    )

    meta = doc.add_paragraph()
    add_run(meta, f"Prepared date: {date(2026, 5, 26).strftime('%B %d, %Y')}\n", bold=True)
    add_run(meta, "Prepared for: Nigerian business legal representative / licensing counsel\n", bold=True)
    add_run(meta, "Prepared by: DARE project team\n", bold=True)
    add_run(meta, "Status: Product and regulatory briefing. This is not legal advice.")

    add_note_box(
        doc,
        "Purpose of this document.",
        "This brief explains what DARE is, how it is intended to operate, where money and user data move, and what legal classification and licensing issues counsel should resolve before incorporation, payment onboarding, or any real-money launch in Nigeria.",
    )

    doc.add_heading("Contents", level=1)
    add_numbered(
        doc,
        [
            "Executive summary",
            "Plain-English product explanation",
            "What DARE is and is not",
            "User roles and user journey",
            "DARE lifecycle and settlement model",
            "Wallet, escrow, fees, and payments",
            "Business formation and operating structure",
            "Nigeria licensing and regulatory questions",
            "Compliance, risk, and responsible play controls",
            "Counsel workplan and documents requested",
            "Appendices: terminology, launch assumptions, references",
        ],
    )

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "DARE is a mobile-first social challenge platform for peer-to-peer skill challenges. A user creates a challenge, another user accepts it, both parties agree to a written rule set called a constitution, funds may be locked in escrow, the challenge runs in a live Court experience, and the result is resolved by answer-key verification, witnessed live result, evidence review, or a dispute review process. The product is designed around rule clarity, escrow transparency, auditability, user trust, and responsible participation."
    )
    doc.add_paragraph(
        "DARE should be reviewed by counsel as a regulated-risk product. Although the commercial positioning is not a generic betting app and the intended first release prioritizes skill-based, answer-key verified contests, the combination of entry value, prize value, peer competition, platform fees, wallets, and payouts means counsel must classify the product under Nigerian state gaming/lottery/prize-competition laws, payment rules, AML/CFT obligations, data protection law, consumer protection rules, advertising standards, tax obligations, and app-store/payment-processor policies."
    )
    doc.add_paragraph(
        "The recommended launch discipline is to begin with one Nigerian jurisdiction, likely Lagos if commercially and legally viable, and only open real-money gameplay after counsel confirms licensing posture, payment-provider approval, KYC/AML operations, responsible gaming controls, and regulatory reporting obligations."
    )

    add_table(
        doc,
        ["Item", "Current DARE position", "Counsel decision needed"],
        [
            [
                "Product category",
                "Skill-first peer challenge platform with optional real-money stakes and escrow.",
                "Confirm whether it is treated as gaming, betting, lottery, prize competition, promotional competition, skill game, or another regulated category in each launch state.",
            ],
            [
                "Launch scope",
                "MVP should support answer_key, witnessed, and evidence DAREs with strict limits, audit trails, and dispute controls before expanding to high-risk physical formats.",
                "Confirm if this narrower MVP changes license class, permitted advertising, payment approval, or consumer terms.",
            ],
            [
                "Money flow",
                "Users deposit funds, stakes are held in platform-controlled escrow ledger, winner receives payout after settlement, platform may take a fee.",
                "Determine whether DARE can hold customer balances directly, must partner with a licensed payment institution, or needs a trust/segregated account structure.",
            ],
            [
                "Responsible play",
                "Age gate, KYC tiers, stake limits, deposit limits, cooling-off, self-exclusion, account freeze, dispute handling, and audit logs are planned.",
                "Confirm mandatory minimum controls and regulator-facing policies.",
            ],
        ],
        [1800, 3780, 3780],
    )

    doc.add_heading("2. Plain-English Product Explanation", level=1)
    doc.add_paragraph(
        "DARE turns informal social challenges into structured, enforceable digital challenges. The everyday behavior is familiar: one person says, 'I dare you to do this,' or 'I bet you cannot beat me at this.' DARE formalizes that behavior by requiring clear rules before anyone commits, locking agreed value where applicable, running the challenge inside an auditable product flow, and releasing funds only after a defined result."
    )
    doc.add_paragraph(
        "A DARE can be thought of as a contract-like challenge object inside the app. Each DARE contains the test, category, duration, proof method, stake, fee, payout preview, settlement method, and edge-case handling. After both sides accept the constitution, the terms become immutable for that DARE unless both parties explicitly accept an amended version."
    )
    doc.add_paragraph(
        "The initial product should avoid open-ended physical dares and uncertain proof workflows. The safer MVP is creator-authored answer_key, witnessed, and evidence flows with clear constitutions, conservative stake limits, and auditable review. Tournaments, creator economy, and wider marketplace features can follow only after licensing, wallet, dispute, and risk operations are proven."
    )

    doc.add_heading("3. What DARE Is and Is Not", level=1)
    add_table(
        doc,
        ["DARE is", "DARE is not intended to be", "Why this distinction matters"],
        [
            [
                "A peer-to-peer skill challenge platform.",
                "A sportsbook where users bet on third-party events.",
                "DARE outcomes are meant to depend on participant action inside a challenge, not passive speculation on external events.",
            ],
            [
                "A structured challenge constitution and settlement system.",
                "An informal social content app with unenforced dares.",
                "The platform creates records, rules, escrow, and settlement expectations that need legal terms and operational controls.",
            ],
            [
                "A wallet, escrow, dispute, and trust infrastructure layer for challenge outcomes.",
                "A bank, payment service bank, or money transfer operator.",
                "Counsel must structure money handling through approved providers, segregated accounts, and clear wallet terms.",
            ],
            [
                "A skill-first entertainment and competition product with responsible play controls.",
                "A casino-style chance product or luck-first gambling interface.",
                "Branding, mechanics, odds, randomization, and advertising should avoid creating unnecessary gaming classification risk.",
            ],
        ],
        [3120, 3120, 3120],
    )

    doc.add_heading("4. User Roles and User Journey", level=1)
    doc.add_heading("Core roles", level=2)
    add_bullets(
        doc,
        [
            "Guest: views limited public marketing and begins registration.",
            "Player: authenticated user who can create, accept, and participate in DAREs once eligibility rules are met.",
            "Issuer: player who creates a DARE and proposes the constitution.",
            "Challenger: player who accepts a DARE after reviewing the constitution.",
            "Spectator: user who may watch live Court sessions and interact where permitted.",
            "Juror: eligible user who reviews blind evidence packets for disputes or evidence-based DAREs.",
            "Admin / risk operator: internal operator who reviews disputes, fraud flags, KYC, payment anomalies, reports, and escalations.",
            "Payment operator: internal or provider-facing operations role handling withdrawals, reconciliation, chargebacks, and payout exceptions.",
        ],
    )
    doc.add_heading("Standard real-money challenge journey", level=2)
    add_numbered(
        doc,
        [
            "User registers, passes age gate, accepts terms, and completes the required KYC tier.",
            "User funds the wallet through an approved payment provider flow.",
            "Issuer creates a DARE with a title, category, stake, duration, rules, proof method, settlement method, and edge-case handling.",
            "Server validates issuer eligibility, KYC tier, risk status, responsible play limits, and wallet balance.",
            "Issuer stake is locked in escrow if the product policy requires issuer lock at creation.",
            "Challenger reviews the constitution, stake, fee, payout preview, cancellation terms, and dispute terms.",
            "Server validates the challenger and locks challenger stake.",
            "Both players enter Court, ready up, and the server starts the match using server time.",
            "The server computes or receives the result according to the DARE type.",
            "Escrow is settled to the winner after the dispute window, unless a valid dispute freezes settlement.",
            "Wallet ledger entries, trust score effects, notifications, and audit logs are recorded.",
        ],
    )

    doc.add_heading("5. DARE Lifecycle and Settlement Model", level=1)
    doc.add_paragraph(
        "The product should be represented to regulators and providers as a state-machine-driven system. Sensitive states are not controlled by the mobile client. The backend owns eligibility checks, money movement, timers, final scoring, dispute windows, settlement, trust score updates, and admin interventions."
    )
    add_table(
        doc,
        ["Lifecycle stage", "Meaning", "Controls"],
        [
            ["Draft", "Issuer is preparing rules and stake.", "No public commitment; no client-side money movement."],
            ["Open / targeted pending", "DARE is available for an opponent or sent to a specific opponent.", "Eligibility checks, expiry, issuer escrow policy, responsible play limits."],
            ["Accepted", "Opponent has accepted terms.", "Both users must have accepted the same constitution version."],
            ["Ready check", "Players prepare for Court.", "Server readiness, heartbeat, reconnect grace, no timer manipulation."],
            ["Active", "Challenge is live.", "Server start/end time, answer-key verified result, anti-cheat, event log."],
            ["Completed", "A provisional result exists.", "Dispute window, evidence preservation, settlement hold if needed."],
            ["Dispute pending / jury open", "A dispute or evidence review is active.", "Escrow remains held; blind evidence packet; juror eligibility; admin escalation."],
            ["Settled", "Escrow released according to final result.", "Append-only ledger, fee entry, payout entry, audit log."],
            ["Cancelled / expired / voided", "DARE did not settle as a win/loss.", "Refund rules and reason codes must be explicit."],
        ],
        [1800, 3840, 3720],
    )

    doc.add_heading("6. Wallet, Escrow, Fees, and Payments", level=1)
    doc.add_paragraph(
        "The wallet is a financial product surface and should be treated as high risk. The app only displays wallet state. It must not decide balances, escrow, payouts, withdrawals, or settlement. DARE's backend records money movement through an append-only ledger and reconciles provider transactions against internal records."
    )
    add_table(
        doc,
        ["Balance type", "Meaning", "Legal / operational issue"],
        [
            ["Available balance", "Funds user can use or request to withdraw.", "Terms must clarify whether this is e-money, stored value, user credit, or provider-held balance."],
            ["Escrowed balance", "Funds locked for active DAREs.", "Counsel should confirm permissible escrow/trust account structure and user rights during disputes."],
            ["Pending balance", "Deposits or withdrawals awaiting provider confirmation.", "Terms and support procedures must explain delays and failed transactions."],
            ["Held balance", "Funds frozen due to dispute, fraud, chargeback, or compliance review.", "Requires clear authority in terms, audit logs, and escalation policy."],
        ],
        [1900, 3600, 3860],
    )
    doc.add_heading("Proposed money flow", level=2)
    add_numbered(
        doc,
        [
            "Deposit: user initiates deposit; provider checkout or bank transfer confirms payment; backend verifies provider reference, amount, currency, and signature; ledger records deposit_confirmed.",
            "Escrow hold: when a DARE is created or accepted, backend checks balance and records escrow_hold entries for each party.",
            "Settlement: after result and dispute window, backend releases escrow, records winner payout, platform fee, and any juror allocation.",
            "Withdrawal: user requests withdrawal; backend checks available balance, KYC tier, limits, risk status, and pending holds; withdrawal is processed through approved provider or operations queue.",
            "Reconciliation: daily jobs compare provider records with internal ledger, detect stuck or duplicate transactions, and report anomalies.",
        ],
    )
    doc.add_heading("Fee questions for counsel", level=2)
    add_bullets(
        doc,
        [
            "Whether the platform fee/rake changes classification under gaming, lottery, prize competition, or betting law.",
            "Whether fees should be deducted at stake entry, settlement, or only from winnings.",
            "Whether tax, withholding, VAT, levy, good-causes levy, gaming levy, or state-level reporting applies.",
            "Whether juror rewards are permitted and how they should be documented and taxed.",
            "Whether user balances must be held in a segregated account or through a licensed payment partner.",
        ],
    )

    doc.add_heading("7. Business Formation and Operating Structure", level=1)
    doc.add_paragraph(
        "Counsel should form a Nigerian operating entity capable of entering payment-provider agreements, applying for state gaming or promotional licenses where required, registering intellectual property, opening corporate bank/provider accounts, contracting with vendors, employing staff, and holding regulator-facing obligations. The exact structure should be selected by counsel after reviewing licensing requirements."
    )
    add_table(
        doc,
        ["Workstream", "What DARE likely needs", "Counsel output requested"],
        [
            [
                "Entity formation",
                "Nigerian company limited by shares, unless counsel recommends another vehicle.",
                "CAC incorporation, constitutional documents, shareholding/director structure, registered address, company objects broad enough for software, digital services, competitions, payments facilitation, and licensed gaming/interactive games if appropriate.",
            ],
            [
                "Trade name and IP",
                "DARE brand, logo, domain, app name, UI assets, product copy, and future class protection.",
                "Trademark search and filing strategy in Nigeria; confirm ownership assignment from developers/designers; prepare IP assignment clauses.",
            ],
            [
                "Tax",
                "TIN, VAT and company tax setup, transfer pricing consideration if foreign parent/affiliate is used.",
                "Tax registration, advisory on fee treatment, player winnings, juror rewards, withholding obligations, and statutory levies.",
            ],
            [
                "Licensing",
                "Possible state gaming/lottery/interactive game/promotional competition license, depending on classification.",
                "Written legal classification memo and licensing route by launch jurisdiction, starting with Lagos if selected.",
            ],
            [
                "Payments",
                "Provider onboarding for deposits, transfers, webhooks, settlement account, and withdrawal operations.",
                "Provider approval pack and review of payment terms, restricted business approvals, account segregation, chargeback policy.",
            ],
            [
                "Data protection",
                "Processing of identity, contact, KYC, wallet, gameplay, risk, device, chat, and evidence media data.",
                "Privacy notice, NDPA compliance plan, data retention policy, DPO/DPCO assessment, breach response, cross-border transfer review.",
            ],
        ],
        [1700, 3760, 3900],
    )

    doc.add_heading("8. Nigeria Licensing and Regulatory Questions", level=1)
    add_note_box(
        doc,
        "Key legal classification issue.",
        "Do not assume that describing DARE as skill-based removes regulatory obligations. Counsel should analyze the full product mechanics: entry value, prize value, peer-to-peer stake, platform fee, wallet balance, chance vs skill, user-generated challenge terms, dispute process, and whether any DARE category introduces chance or public lottery-like characteristics.",
    )
    doc.add_heading("Gaming, lottery, betting, and prize competition", level=2)
    doc.add_paragraph(
        "Nigeria's gaming/licensing position requires state-by-state review. Public reporting on the Supreme Court's November 22, 2024 decision indicates that the National Lottery Act 2005 should no longer apply in all states except the Federal Capital Territory for lottery/games-of-chance regulation, with state Houses of Assembly having the relevant authority. Lagos State's regulator states that it regulates and issues retail, virtual, and online operating licenses across categories including public online lottery, online sports betting, scratch cards, interactive games, casino, gaming machine, pool betting, promotional competitions, remote gaming, and online/retail gaming related activities."
    )
    doc.add_paragraph(
        "Because DARE includes entry value and prize-like payout potential, counsel should confirm whether Lagos or any other target state would classify DARE under interactive games, remote gaming, online gaming, promotional competition, other games, betting, lottery, or a separate skill competition category."
    )
    doc.add_heading("Payment provider policy", level=2)
    doc.add_paragraph(
        "Payment approval is separate from statutory legality. Paystack's public terms restrict gambling, gaming, and other entry-fee/prize activities, including games of skill, unless prior approval is obtained and the operator and customers are in jurisdictions where the activity is permitted by law. Counsel should prepare a provider-facing product description and approval memo before deposits or payouts are enabled."
    )
    doc.add_heading("Payments and stored value", level=2)
    doc.add_paragraph(
        "Counsel should confirm whether DARE is merely a merchant using a licensed payment provider, a marketplace with split settlement, a wallet/stored-value product, an escrow service, a payment facilitator, or another regulated payment arrangement. The Central Bank of Nigeria maintains payment-system regulations and license categories, including payment service provider categorization, electronic payment channels, instant inter-bank transfers, QR, mobile money, payment service banks, and other frameworks."
    )
    doc.add_heading("AML/CFT and suspicious activity", level=2)
    doc.add_paragraph(
        "DARE can be abused for collusion, money movement, account farming, chargeback fraud, and staged losses. If classified within casino/gaming/pool betting/lottery or another designated sector, SCUML/EFCC registration and AML/CFT controls may be required. Even if a provider handles payment rails, DARE still needs KYC tiers, transaction monitoring, suspicious activity escalation, records, sanctions screening where appropriate, and limits."
    )
    doc.add_heading("Data protection", level=2)
    doc.add_paragraph(
        "DARE will process personal data and potentially sensitive identity/evidence data: phone/email, profile data, KYC submissions, device/session data, payment metadata, chat, disputes, evidence media, risk flags, and admin notes. Counsel should determine NDPA controller/processor obligations, registration/audit obligations, privacy notice requirements, lawful bases, retention periods, cross-border transfers, breach reporting, data subject rights, and vendor DPAs."
    )
    doc.add_heading("Consumer protection, advertising, and responsible play", level=2)
    add_bullets(
        doc,
        [
            "Age gate and underage prevention requirements.",
            "Responsible gaming/play language, deposit limits, stake limits, timeouts, cooling-off, and self-exclusion.",
            "Marketing restrictions: avoid misleading income claims, guaranteed winnings, casino-style copy, pressure tactics, or targeting minors.",
            "Clear terms: challenge rules, cancellation, refund, voiding, disputes, wallet holds, platform fees, payout timing, support channels, and account sanctions.",
            "Complaint handling and dispute escalation to regulator or internal ombuds process if required.",
        ],
    )

    doc.add_heading("9. Compliance, Risk, and Responsible Play Controls", level=1)
    add_table(
        doc,
        ["Risk", "Example abuse", "Planned control"],
        [
            ["Underage participation", "Minor registers and plays for money.", "Age gate, KYC, document checks, account restrictions, parental/minor reporting workflow."],
            ["Money laundering / collusion", "Users stage repeated predictable outcomes to move funds.", "Velocity limits, repeated matchup detection, shared device/IP signals, source-of-funds checks, suspicious activity review."],
            ["Payment fraud", "Chargeback after loss or duplicate deposit credit.", "Webhook signature verification, provider reference uniqueness, idempotency keys, reconciliation, chargeback hold."],
            ["Client manipulation", "Client attempts to forge score, timer, wallet balance, or winner.", "Server-authoritative state machine, server time, row locks, audited RPCs, no client-side settlement."],
            ["Bad-faith disputes", "Losing party disputes every result to delay payout.", "Dispute window, rationale requirement, trust penalties, repeat dispute limits, admin review."],
            ["Jury capture", "Participant coordinates with jurors.", "Random server-side assignment, blind packets, relationship checks, hidden vote tallies, immutable votes."],
            ["Evidence manipulation", "Edited or substituted media.", "In-app capture where required, content hash, private storage, signed URLs, upload confirmation, evidence access logs."],
            ["Problem play", "User repeatedly stakes beyond healthy limits.", "Deposit/stake limits, cooling-off, self-exclusion, responsible play notices, admin limits."],
        ],
        [1800, 3460, 4100],
    )
    doc.add_heading("Launch blockers for real money", level=2)
    add_bullets(
        doc,
        [
            "Written legal classification memo for the launch jurisdiction.",
            "Confirmed license path or written legal basis for operating without a specific gaming license.",
            "Payment provider approval for skill-based entry-fee/prize activities.",
            "KYC/AML policy and operational owner.",
            "Responsible play policy and self-exclusion process.",
            "Data protection documentation, including privacy notice and breach process.",
            "Audited wallet ledger, escrow, idempotency, and reconciliation tests.",
            "Admin console with audit logs, freeze controls, and dispute review.",
            "Approved user terms, privacy policy, wallet terms, challenge constitution terms, and dispute rules.",
        ],
    )

    doc.add_heading("10. Counsel Workplan and Documents Requested", level=1)
    doc.add_heading("A. Immediate legal analysis", level=2)
    add_numbered(
        doc,
        [
            "Classify the MVP answer_key DARE product under Nigerian law for the proposed launch state.",
            "Identify whether the platform needs a Lagos State Lotteries and Gaming Authority license or equivalent state license before beta or public launch.",
            "Identify whether federal/FCT licenses, state registrations, or inter-state restrictions apply if users from multiple states can access the app.",
            "Determine whether product mechanics must change to fit a lower-risk category, such as capped-stake skill competitions, no-rake competitions, promotional competitions, or non-cash rewards.",
            "Confirm minimum age, responsible play, advertising, complaint, data retention, and user fund segregation obligations.",
            "Confirm payment provider approval path and whether provider terms allow deposits, wallet balances, escrow, and payouts.",
        ],
    )
    doc.add_heading("B. Formation and regulatory setup", level=2)
    add_bullets(
        doc,
        [
            "Incorporate the operating company with suitable objects.",
            "Register tax ID and advise on VAT/company tax/levy treatment.",
            "Prepare BO/PSC records, board resolutions, officer appointments, company secretary requirements, and annual compliance calendar.",
            "File trademark applications for DARE wordmark/logo and review domain/social/app-store name conflicts.",
            "Prepare license application package for the selected regulator, including business plan, product description, responsible gaming policy, AML/CFT policy, technical controls, payment flow, and financial projections if required.",
            "Prepare payment provider onboarding and restricted-business approval package.",
        ],
    )
    doc.add_heading("C. Product legal documents", level=2)
    add_bullets(
        doc,
        [
            "Terms of service with challenge constitution incorporation, user obligations, prohibited conduct, dispute rules, wallet rules, fee disclosures, refund/void rules, sanctions, and limitation of liability.",
            "Privacy notice under NDPA, including KYC, payments, evidence, chat, risk profiling, retention, cross-border processing, and user rights.",
            "Responsible play policy with age restrictions, limits, self-exclusion, cooling-off, and support information.",
            "AML/KYC policy for internal operations and provider/regulator review.",
            "Dispute and jury policy explaining evidence submission, juror eligibility, vote finality, admin escalation, and escrow holds.",
            "Community rules, acceptable content policy, anti-harassment policy, and sanctions policy.",
            "Payment/wallet terms explaining deposits, withdrawals, pending funds, held funds, chargebacks, account freezes, provider downtime, and reconciliation.",
        ],
    )
    doc.add_heading("D. Questions counsel should answer in writing", level=2)
    add_bullets(
        doc,
        [
            "Can DARE lawfully launch answer_key real-money DAREs in Lagos without being licensed as betting/gaming/lottery/interactive gaming? If not, what license class applies?",
            "Does the platform fee create or increase gambling/gaming classification risk?",
            "Are peer-to-peer stakes treated differently from platform-funded prizes?",
            "Can DARE hold user funds internally, or must all funds remain with a licensed payment provider/trust account?",
            "What state-by-state geofencing is required if only one state license is obtained?",
            "What are the required regulator disclosures, reports, technical certifications, and levies?",
            "What KYC tier is required before deposit, before accepting a DARE, before withdrawal, and before high-value stakes?",
            "Can jurors receive rewards, and how should those payments be classified?",
            "What evidence retention period is required, and when must evidence be deleted?",
            "What language must be used or avoided in marketing to reduce regulatory and consumer-protection risk?",
        ],
    )

    doc.add_heading("11. Appendix A: Proposed MVP Boundaries", level=1)
    add_table(
        doc,
        ["Included in first real-money MVP", "Deferred until after licensing/operations validation"],
        [
            [
                "Answer Key, Witnessed, and Evidence DAREs; wallet deposit; escrow hold; ready-up Court; server timer; result screen; settlement; dispute filing; admin queue; KYC tiering; limits; audit logs.",
                "High-risk physical formats; unsupported handshake-only outcomes with real stakes; tournaments; creator economy; spectator rewards; multi-country launch; USSD production gateway; AI voice-to-DARE.",
            ]
        ],
        [4680, 4680],
    )

    doc.add_heading("12. Appendix B: Key Terms", level=1)
    add_table(
        doc,
        ["Term", "Meaning in DARE"],
        [
            ["DARE", "A structured challenge between users, governed by a constitution and resolved through a defined method."],
            ["Constitution", "The binding rule set for a DARE: test, duration, stake, proof method, edge cases, and settlement terms."],
            ["Court", "The live match interface where participants ready up, play, submit proof, or complete the challenge."],
            ["Escrow", "Internal hold or provider-backed segregation of funds pending challenge result or dispute resolution."],
            ["Juror", "Eligible user assigned to review blind evidence and vote on disputes or evidence-based outcomes."],
            ["Trust score", "Server-calculated reputation metric based on completed DAREs, disputes, forfeits, jury behavior, and risk events."],
            ["Answer Key DARE", "DARE resolved by answer-key verified rules, such as timed quiz or deterministic answers."],
            ["Evidence DARE", "DARE resolved using submitted media or proof, usually with jury/admin review."],
        ],
        [2100, 7260],
    )

    doc.add_heading("13. Appendix C: Product and Regulatory References Checked", level=1)
    doc.add_paragraph(
        "The following references were reviewed for this business briefing. Counsel should treat them as starting points and perform independent legal verification before filing or advising."
    )
    refs = [
        ("DARE internal product brief", "docs/01-product-brief.md"),
        ("DARE internal wallet, escrow, and payments brief", "docs/06-wallet-escrow-and-payments.md"),
        ("DARE internal disputes, jury, and trust brief", "docs/07-disputes-jury-and-trust.md"),
        ("DARE internal security, risk, and compliance brief", "docs/08-security-risk-and-compliance.md"),
        ("Lagos State Lotteries and Gaming Authority - regulatory role", "https://lslga.org/about/"),
        ("Central Bank of Nigeria - Payment Service Providers and payment system rules", "https://www.cbn.gov.ng/PaymentsSystem/PSPs.html"),
        ("Nigeria Data Protection Commission - Nigeria Data Protection Act 2023", "https://www.ndpc.gov.ng/ndp-act-2023/"),
        ("Paystack terms - acceptable use restriction for gaming/games of skill with entry fee and prize", "https://paystack.com.ng/terms"),
        ("SCUML/EFCC registration guidance categories", "https://scumlportal.efcc.gov.ng/scuml-registration-guideline.php"),
        ("NFIU laws and regulations page", "https://nfiu.gov.ng/LawsAndRegulations"),
        ("Corporate Affairs Commission", "https://www.cac.gov.ng/"),
        ("Nigeria Startup Portal", "https://www.startup.gov.ng/"),
        ("Channels Television report on Supreme Court National Lottery Act decision, November 22, 2024", "https://www.channelstv.com/2024/11/22/supreme-court-nullifies-national-lottery-act/"),
    ]
    for label, url in refs:
        p = doc.add_paragraph(style="List Bullet")
        if url.startswith("http"):
            add_hyperlink(p, label, url)
            add_run(p, f" ({url})")
        else:
            add_run(p, label + ": ")
            add_run(p, url)

    doc.add_heading("14. Closing Instruction to Counsel", level=1)
    doc.add_paragraph(
        "Please review DARE as a financial-risk, gaming-classification, data-protection, and consumer-protection product before incorporation and launch. The project team needs a written go/no-go licensing path, recommended company structure, required regulator approvals, required payment-provider approvals, mandatory policy documents, and any product changes required to lawfully run a controlled Nigerian beta."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
